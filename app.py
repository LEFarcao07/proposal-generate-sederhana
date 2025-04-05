import os
from flask import Flask, request, send_file, render_template, after_this_request, send_from_directory, jsonify
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

app = Flask(__name__, template_folder=os.getcwd())

# Enable CORS
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
    response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
    return response

# Path untuk folder assets dan temp
ASSETS_DIR = os.path.join(os.getcwd(), 'assets')
TEMP_DIR = os.path.join(ASSETS_DIR, 'temp')

os.makedirs(ASSETS_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory(os.getcwd(), filename)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST', 'OPTIONS'])
def generate_docx():
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'}), 200
        
    try:
        if 'logo' not in request.files:
            return jsonify({"error": "No file part"}), 400
            
        logo = request.files['logo']
        
        if logo.filename == '':
            return jsonify({"error": "No selected file"}), 400
            
        if not logo or not allowed_file(logo.filename):
            return jsonify({"error": "Format file tidak didukung. Harap unggah file dengan format: PNG, JPG, atau JPEG."}), 400

        logo_path = os.path.join(ASSETS_DIR, 'logo.' + logo.filename.rsplit('.', 1)[1].lower())
        logo.save(logo_path)

        form_data = {
            'judul': request.form.get('judul', ''),
            'subjudul': request.form.get('subjudul', ''),
            'nama_kelompok': request.form.get('nama_kelompok', ''),
            'informasi_lain': request.form.get('informasi_lain', ''),
            'kata_pengantar': request.form.get('kata_pengantar', ''),
            'latar_belakang': request.form.get('latar_belakang', ''),
            'visi': request.form.get('visi', ''),
            'misi': request.form.get('misi', ''),
            'tujuan': request.form.get('tujuan', '')
        }

        for field, value in form_data.items():
            if not value.strip():
                return jsonify({"error": f"Field {field} tidak boleh kosong"}), 400

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Judul Proposal - CENTERED
        add_html_paragraph(doc, form_data['judul'], alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(28))

        # Subjudul Proposal - CENTERED
        add_html_paragraph(doc, form_data['subjudul'], alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(17))

        # Spacing
        doc.add_paragraph().paragraph_format.space_after = Pt(34)

        # Logo - CENTERED
        try:
            logo_width = Inches(2.5)
            doc.add_picture(logo_path, width=logo_width)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error adding logo: {str(e)}")
            return jsonify({"error": "Gagal memproses gambar logo. Pastikan file gambar valid."}), 400

        doc.add_paragraph().paragraph_format.space_after = Pt(16)

        # Nama/Kelompok - CENTERED
        add_html_paragraph(doc, form_data['nama_kelompok'], alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                         font_size=Pt(14), line_spacing=1.0, paragraph_spacing=Pt(12), 
                         prefix="Oleh :")

        doc.add_paragraph().paragraph_format.space_after = Pt(108)

        # Informasi Lain - CENTERED
        add_html_paragraph(doc, form_data['informasi_lain'], alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                         font_size=Pt(12), line_spacing=1.0, paragraph_spacing=Pt(12))
        
        doc.add_page_break()
        
        daftar_isi_title = doc.add_paragraph('DAFTAR ISI')
        daftar_isi_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in daftar_isi_title.runs:
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        def add_daftar_isi_line(doc, text):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            paragraph.add_run(text).bold = False
            paragraph.add_run("\t")
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            run._r.append(instrText)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
                
        add_daftar_isi_line(doc, 'SAMPUL')
        add_daftar_isi_line(doc, 'DAFTAR ISI')
        add_daftar_isi_line(doc, 'KATA PENGANTAR')
        add_daftar_isi_line(doc, 'BAB I')
        add_daftar_isi_line(doc, '   LATAR BELAKANG')
        add_daftar_isi_line(doc, '   1.1 LATAR BELAKANG')
        add_daftar_isi_line(doc, 'BAB II')
        add_daftar_isi_line(doc, '   VISI, MISI, DAN TUJUAN')
        add_daftar_isi_line(doc, '   2.1 VISI')
        add_daftar_isi_line(doc, '   2.2 MISI')
        add_daftar_isi_line(doc, '   2.3 TUJUAN')
        
        doc.add_page_break()

        # Kata Pengantar
        kata_pengantar_title = doc.add_paragraph('KATA PENGANTAR')
        kata_pengantar_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in kata_pengantar_title.runs:
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        add_html_paragraph(doc, form_data['kata_pengantar'], font_size=Pt(12))

        doc.add_page_break()

        # BAB I: Latar Belakang Perusahaan (Title)
        bab1_title = doc.add_paragraph('BAB I: LATAR BELAKANG')
        bab1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in bab1_title.runs:
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Latar (Subtitle)
        latar_belakang_title = doc.add_paragraph('1.1 Latar Belakang')
        latar_belakang_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in latar_belakang_title.runs:
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Isi Latar
        add_html_paragraph(doc, form_data['latar_belakang'], font_size=Pt(12))

        # Halaman Baru (BAB II)
        doc.add_page_break()

        # BAB II: Visi, Misi, dan Tujuan Perusahaan (Title)
        bab2_title = doc.add_paragraph('BAB II: VISI, MISI, DAN TUJUAN')
        bab2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in bab2_title.runs:
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Visi (Subtitle)
        visi_title = doc.add_paragraph('2.1 Visi')
        visi_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in visi_title.runs:
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Isi Visi
        add_html_paragraph(doc, form_data['visi'], font_size=Pt(12))

        # Tambahkan jarak vertikal
        doc.add_paragraph().paragraph_format.space_after = Pt(24)

        # Misi (Subtitle)
        misi_title = doc.add_paragraph('2.2 Misi')
        misi_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in misi_title.runs:
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Isi Misi
        add_html_paragraph(doc, form_data['misi'], font_size=Pt(12))

        # Tambahkan jarak vertikal
        doc.add_paragraph().paragraph_format.space_after = Pt(24)

        # Tujuan (Subtitle)
        tujuan_title = doc.add_paragraph('2.3 Tujuan')
        tujuan_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in tujuan_title.runs:
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Isi Tujuan
        add_html_paragraph(doc, form_data['tujuan'], font_size=Pt(12))
        
        def add_page_number(doc):
            section = doc.sections[0]
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            run._r.append(instrText)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)

        add_page_number(doc)

        # Save and return the document
        doc_path = os.path.join(TEMP_DIR, 'proposal.docx')
        doc.save(doc_path)

        @after_this_request
        def remove_file(response):
            try:
                if os.path.exists(doc_path):
                    os.remove(doc_path)
                if os.path.exists(logo_path):
                    os.remove(logo_path)
            except Exception as e:
                print(f"Gagal menghapus file sementara: {e}")
            return response

        return send_file(
            doc_path,
            as_attachment=True,
            download_name='proposal.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({"error": f"Terjadi kesalahan: {str(e)}"}), 500

def add_html_paragraph(doc, html, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(12), 
                      prefix=None, line_spacing=1.0, paragraph_spacing=Pt(12)):
    if not html or html.strip() == '':
        return

    soup = BeautifulSoup(html, 'html.parser')
    
    if prefix:
        prefix_paragraph = doc.add_paragraph(prefix)
        prefix_paragraph.alignment = alignment
        prefix_paragraph.paragraph_format.space_after = paragraph_spacing
        for run in prefix_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = font_size

    def process_element(element, paragraph=None, current_formats=None):
        if current_formats is None:
            current_formats = {
                'bold': False,
                'italic': False,
                'underline': False
            }
            
        if isinstance(element, str):
            if element.strip():
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = alignment
                    paragraph.paragraph_format.line_spacing = line_spacing
                    paragraph.paragraph_format.space_after = paragraph_spacing
                
                run = paragraph.add_run(element)
                run.font.name = 'Times New Roman'
                run.font.size = font_size
                run.bold = current_formats['bold']
                run.italic = current_formats['italic']
                run.underline = current_formats['underline']
            return
        
        # Handle different HTML elements
        if element.name in ['p', 'div']:
            paragraph = doc.add_paragraph()
            paragraph.alignment = alignment
            paragraph.paragraph_format.line_spacing = line_spacing
            paragraph.paragraph_format.space_after = paragraph_spacing
            
            for child in element.children:
                process_element(child, paragraph, current_formats.copy())
                
        elif element.name in ['b', 'strong']:
            new_formats = current_formats.copy()
            new_formats['bold'] = True
            for child in element.children:
                process_element(child, paragraph, new_formats)
                
        elif element.name in ['i', 'em']:
            new_formats = current_formats.copy()
            new_formats['italic'] = True
            for child in element.children:
                process_element(child, paragraph, new_formats)
                
        elif element.name == 'u':
            new_formats = current_formats.copy()
            new_formats['underline'] = True
            for child in element.children:
                process_element(child, paragraph, new_formats)
                
        elif element.name == 'br':
            if paragraph:
                paragraph.add_run().add_break()
                
        elif element.name in ['ul', 'ol']:
            list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
            for li in element.find_all('li', recursive=False):
                paragraph = doc.add_paragraph(style=list_style)
                paragraph.alignment = alignment
                paragraph.paragraph_format.line_spacing = line_spacing
                paragraph.paragraph_format.space_after = paragraph_spacing
                
                for child in li.children:
                    process_element(child, paragraph, current_formats.copy())
                    
        elif element.name == 'li':
            for child in element.children:
                process_element(child, paragraph, current_formats.copy())

    # Process each top-level element
    for element in soup.children:
        process_element(element)

if __name__ == '__main__':
    app.run(debug=True, port=3020, host='0.0.0.0')